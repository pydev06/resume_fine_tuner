from docx import Document
import io
import difflib

def generate_tailored_docx(file_bytes: bytes, rewrites: list) -> bytes:
    """
    Load a DOCX file, apply rewrites using fuzzy matching, and return the modified file bytes.
    PRESERVES original formatting by modifying runs in-place where possible.
    """
    if not rewrites:
        return file_bytes

    doc = Document(io.BytesIO(file_bytes))
    
    # Pre-process rewrites for easier access
    # We'll use a list of objects so we can track if they've been used
    pending_rewrites = [{"original": r["original"].strip(), "rewritten": r["rewritten"]} for r in rewrites if "original" in r]

    # Iterate through all paragraphs in the document
    # Note: This doesn't cover tables yet, but covers most resume bullets
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
            
        clean_text = para.text.strip()
        
        # Check against pending rewrites
        match_index = -1
        highest_ratio = 0.0
        
        for idx, item in enumerate(pending_rewrites):
            ratio = difflib.SequenceMatcher(None, clean_text, item["original"]).ratio()
            if ratio > highest_ratio:
                highest_ratio = ratio
                match_index = idx
        
        # Threshold for replacement
        if highest_ratio > 0.85:
            rewrite_item = pending_rewrites[match_index]
            new_text = rewrite_item["rewritten"]
            
            # Found a match! Now replace text while trying to preserve formatting.
            print(f"DOCX Match: '{clean_text[:30]}...' -> Rewritten ({highest_ratio:.2f})")
            
            # Naive strategy: Replace the text of the first run, clear others
            # Better strategy: if single run, replace. if multiple, it's complex.
            # For resumes, bullets are often single runs or simple styles.
            
            if para.runs:
                # Set first run to new text
                para.runs[0].text = new_text
                # Clear subsequent runs to avoid "ghost" text
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(new_text)

    # Save to buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer.getvalue()
